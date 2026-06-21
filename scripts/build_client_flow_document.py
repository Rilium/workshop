from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "Flusso_Cliente_FunniFin.docx"
OUT_MD = ROOT / "docs" / "flusso-cliente-dettagli.md"

CLIENT_FLOW = ROOT / "manual-screenshots" / "client-flow"
MANUAL = ROOT / "manual-screenshots" / "annotated"


STEPS = [
    {
        "title": "1. Interessi",
        "screen": "01-cliente-interessi.png",
        "manual": ["fig-02-cliente-partenza-vuota.png", "fig-04-cliente-interessi-temi.png", "fig-05-cliente-filtri-temi.png"],
        "does": [
            "Il cliente entra nella vista pubblica, senza autenticazione e senza elementi riservati.",
            "Seleziona uno o piu interessi; i temi collegati vengono attivati e possono essere raffinati.",
            "La CTA resta bloccata finche non esiste almeno una scelta utile per generare il percorso.",
        ],
        "checks": [
            "Stepper visibile sul percorso cliente.",
            "Bottom bar con stato, totale e hint di avanzamento.",
            "Accesso area riservata separato dal percorso pubblico.",
        ],
    },
    {
        "title": "2. Consigliati",
        "screen": "02-cliente-consigliati.png",
        "manual": ["fig-03-cliente-stepper-e-stato.png", "fig-06-cliente-workshop-preventivo.png"],
        "does": [
            "Il sistema propone una prima combinazione di workshop coerente con interessi e temi scelti.",
            "Il cliente puo aggiungere tutti i consigli con un clic oppure saltare al catalogo manuale.",
            "I workshop consigliati non entrano nel carrello finche il cliente non conferma.",
        ],
        "checks": [
            "Numero interessi e numero consigli esplicitati.",
            "Card con prezzo, durata, formato e azione Aggiungi.",
            "Azione secondaria Scegli manualmente disponibile nel footer.",
        ],
    },
    {
        "title": "3. Workshop",
        "screen": "03-cliente-workshop.png",
        "manual": ["fig-06-cliente-workshop-preventivo.png", "fig-07-cliente-catalogo-card.png", "fig-08-cliente-prezzo-sconti.png"],
        "does": [
            "Il cliente rivede catalogo, carrello e preventivo aggiornato in tempo reale.",
            "Puo cercare workshop, aprire filtri per interesse/tema/formato e rimuovere o aggiungere elementi.",
            "Le scelte alimentano riepilogo e-commerce, sconti e CTA verso la personalizzazione.",
        ],
        "checks": [
            "Catalogo filtrato dagli interessi, salvo ricerca o apertura catalogo completo.",
            "Carrello laterale con totale percorso.",
            "Bottom action bar sempre raggiungibile.",
        ],
    },
    {
        "title": "4. Personalizza",
        "screen": "04-cliente-personalizza.png",
        "manual": ["fig-06-cliente-workshop-preventivo.png", "fig-11-bottom-sheet-fisso.png"],
        "does": [
            "Il cliente decide se trasformare ogni workshop in una versione su misura.",
            "La personalizzazione aggiunge co-design FunniFin, note e costo extra per il singolo workshop.",
            "La rimozione dal percorso resta disponibile anche in questa fase.",
        ],
        "checks": [
            "Lista dei workshop selezionati.",
            "Toggle Rendi su misura per ogni riga.",
            "Preventivo aggiornabile in tempo reale.",
        ],
    },
    {
        "title": "5. Date",
        "screen": "05-cliente-date.png",
        "manual": ["fig-09-cliente-date-picker.png", "fig-10-cliente-date-picker-closeup.png"],
        "does": [
            "Il cliente propone una data per ogni workshop selezionato.",
            "Il selettore calendario permette scelta rapida con Adesso, navigazione mese e conferma proposta.",
            "La CTA Carica materiali si abilita solo quando tutti i workshop hanno una data confermata.",
        ],
        "checks": [
            "Stato data per ogni workshop: non scelta, proposta o confermata.",
            "Data, orario e durata visibili nella card dopo conferma.",
            "Hint nel footer se manca almeno una data.",
        ],
    },
    {
        "title": "6. Materiali",
        "screen": "06-cliente-materiali.png",
        "manual": ["fig-11-bottom-sheet-fisso.png", "fig-12-bottom-sheet-disabled.png"],
        "does": [
            "Il cliente puo allegare logo, brand guideline e note platea.",
            "L'upload prepara una cartella Drive draft intestata al cliente; se il flusso viene abbandonato prima dell'invio, la cartella draft viene cestinata.",
            "Lo step e opzionale: il cliente puo procedere all'invio anche senza materiali.",
        ],
        "checks": [
            "Comando Carica materiali visibile.",
            "Messaggio sul comportamento della cartella draft.",
            "Azione Vai all'invio disponibile nel footer.",
        ],
    },
    {
        "title": "7. Invio",
        "screen": "07-cliente-invio.png",
        "manual": ["fig-25-email-richiesta-cliente.png", "fig-26-email-workflow.png"],
        "does": [
            "Il cliente inserisce i dati di contatto solo alla fine del percorso.",
            "Il pannello di readiness ricapitola workshop selezionati e segnala eventuali date mancanti.",
            "L'invio salva la richiesta e prepara il recap email per cliente e FunniFin.",
        ],
        "checks": [
            "Campi obbligatori: nome, cognome, email aziendale, azienda, telefono.",
            "Preventivo finale e materiali collegati vengono associati alla richiesta.",
            "Se l'email non parte, la richiesta resta comunque tracciata come salvata.",
        ],
    },
]


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_picture(doc, path, caption, max_width=6.2, max_height=7.2):
    image_path = Path(path)
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    ratio = height_px / width_px
    width = max_width
    height = width * ratio
    if height > max_height:
        height = max_height
        width = height / ratio

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(width), height=Inches(height))
    add_caption(doc, caption)


def add_bullets(doc, title, items):
    doc.add_paragraph(title, style="Heading 3")
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_docx():
    doc = Document()
    set_doc_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Flusso Cliente FunniFin")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("123832")

    subtitle = doc.add_paragraph()
    subtitle.add_run("Allegato separato al manuale utente - solo vista Cliente").italic = True

    intro = doc.add_paragraph(
        "Questo documento isola l'intero percorso pubblico del cliente: dalla scelta degli interessi all'invio della richiesta. "
        "Le figure principali sono nuove catture dello stesso flusso; le figure di riferimento riprendono le immagini annotate gia usate nel manuale utente."
    )
    intro.paragraph_format.space_after = Pt(12)

    doc.add_paragraph("Perimetro", style="Heading 1")
    for item in [
        "Include esclusivamente il ruolo Cliente e la vista pubblica.",
        "Esclude login, console FunniFin, esperto, brand, notifiche riservate e backend operativo.",
        "Mantiene i riferimenti del manuale solo quando servono a spiegare componenti condivisi del flusso cliente.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Sequenza completa", style="Heading 1")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Step"
    headers[1].text = "Obiettivo cliente"
    headers[2].text = "Uscita attesa"
    overview = [
        ("Interessi", "Selezionare ambiti e temi", "Consigli o catalogo abilitati"),
        ("Consigliati", "Valutare proposta automatica", "Workshop consigliati aggiunti o salto manuale"),
        ("Workshop", "Comporre il carrello", "Percorso e preventivo aggiornati"),
        ("Personalizza", "Richiedere su misura", "Opzioni custom applicate o saltate"),
        ("Date", "Proporre calendario", "Date confermate per tutti i workshop"),
        ("Materiali", "Caricare asset opzionali", "Materiali collegati o step saltato"),
        ("Invio", "Inserire contatti e inviare", "Richiesta salvata e recap preparato"),
    ]
    for row in overview:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    for index, step in enumerate(STEPS):
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_paragraph(step["title"], style="Heading 1")
        add_picture(doc, CLIENT_FLOW / step["screen"], f"Schermata flusso cliente - {step['title']}")
        add_bullets(doc, "Cosa succede", step["does"])
        add_bullets(doc, "Controlli da verificare", step["checks"])
        doc.add_paragraph("Immagini riprese dal manuale utente", style="Heading 2")
        for manual_image in step["manual"]:
            add_picture(doc, MANUAL / manual_image, f"Riferimento manuale utente: {manual_image}", max_width=5.6, max_height=4.2)
        if index != len(STEPS) - 1:
            doc.add_paragraph()

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def rel(path):
    return path.relative_to(ROOT).as_posix()


def build_markdown():
    lines = [
        "# Flusso Cliente FunniFin",
        "",
        "Allegato separato al manuale utente. Perimetro: solo vista Cliente, senza login, console FunniFin, esperto, brand o backend operativo.",
        "",
        "## Sequenza",
        "",
    ]
    for step in STEPS:
        lines.extend([
            f"## {step['title']}",
            "",
            f"![{step['title']}]({rel(CLIENT_FLOW / step['screen'])})",
            "",
            "### Cosa succede",
            "",
        ])
        lines.extend([f"- {item}" for item in step["does"]])
        lines.extend(["", "### Controlli", ""])
        lines.extend([f"- {item}" for item in step["checks"]])
        lines.extend(["", "### Immagini riprese dal manuale utente", ""])
        for manual_image in step["manual"]:
            lines.append(f"![{manual_image}]({rel(MANUAL / manual_image)})")
            lines.append("")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    print(OUT_DOCX)
    print(OUT_MD)
